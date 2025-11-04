import streamlit as st
import pandas as pd
from PyPDF2 import PdfReader
from groq import Groq
import os
import re
from textwrap import wrap
from dotenv import load_dotenv
from docx import Document
from fpdf import FPDF
import tempfile

# ------------------------------
# Load Environment Variables
# ------------------------------
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
LLM_MODEL = os.getenv("LLM_MODEL", "meta-llama/llama-3.3-70b-versatile")
CHAT_MODEL = os.getenv("CHAT_MODEL", "groq/compound")

if not GROQ_API_KEY:
    st.error("⚠️ GROQ_API_KEY not found. Add it in your .env file.")
    st.stop()

client = Groq(api_key=GROQ_API_KEY)


# ------------------------------
# PDF Processing
# ------------------------------
def read_pdf(file):
    """Extracts legally relevant text from a PDF."""
    try:
        reader = PdfReader(file)
        relevant_text = []
        keywords = [
            "ipc section", "indian penal code", "constitution of india",
            "article", "fundamental right", "fundamental duty",
            "preamble", "schedule", "act", "legal", "law", "court", "justice"
        ]

        for page in reader.pages:
            text = page.extract_text() or ""
            text_lower = text.lower()
            if any(kw in text_lower for kw in keywords):
                relevant_text.append(text.strip())

        return "\n\n".join(relevant_text)
    except Exception as e:
        return f"❌ Error reading PDF: {e}"


# ------------------------------
# Legal Relevance Detection
# ------------------------------
def is_law_related(text):
    """Detects legal context in text."""
    text_lower = text.lower()
    keywords = [
        "ipc section", "indian penal code", "constitution of india",
        "fundamental right", "fundamental duty", "article",
        "act", "legal", "law", "court", "justice"
    ]
    return any(kw in text_lower for kw in keywords) or bool(re.search(r"\bsection \d+\b", text_lower))


# ------------------------------
# Text Summarization (Groq)
# ------------------------------
def generate_summary(long_text):
    """Summarizes long legal text in chunks using Groq."""
    try:
        # Split text into chunks of ~4000 characters
        chunks = wrap(long_text, 4000)
        summaries = []

        for chunk in chunks:
            response = client.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {"role": "system", "content": "You are a senior legal expert. Summarize this content clearly."},
                    {"role": "user", "content": chunk}
                ],
                max_tokens=1200
            )
            summaries.append(response.choices[0].message.content.strip())

        merged_text = "\n\n".join(summaries)

        final_response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": "Combine all partial summaries into one well-structured summary."},
                {"role": "user", "content": merged_text}
            ],
            max_tokens=1500
        )

        return final_response.choices[0].message.content.strip()

    except Exception as e:
        msg = str(e).lower()
        if "rate_limit" in msg:
            return "⚠️ Groq API rate limit reached. Please try again later."
        elif "invalid_api_key" in msg:
            return "❌ Invalid Groq API key. Check your .env file."
        elif "model_decommissioned" in msg:
            return "❌ Model decommissioned. Update to a supported model."
        else:
            return f"❌ Summary generation failed: {e}"


# ------------------------------
# Load Constitution Dataset
# ------------------------------
@st.cache_data
def load_coi_dataset():
    """Loads the Constitution of India dataset."""
    try:
        project_root = os.path.dirname(os.path.dirname(__file__))
        csv_path = os.path.join(project_root, "dataset", "coi.csv")
        if not os.path.exists(csv_path):
            raise FileNotFoundError(f"Dataset not found at {csv_path}")
        return pd.read_csv(csv_path)
    except Exception as e:
        st.error(f"⚠️ Error loading COI dataset: {e}")
        return pd.DataFrame()


# ------------------------------
# Export Chat to File
# ------------------------------
def export_chat(format_type="pdf"):
    """Exports chat history as PDF or DOCX."""
    messages = st.session_state.get("law_messages", [])
    if not messages:
        st.warning("⚠️ No chat messages to export.")
        return

    if format_type == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for msg in messages:
            pdf.multi_cell(0, 10, f"{msg['role'].capitalize()}: {msg['content']}")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            pdf.output(tmp.name)
            st.download_button("📄 Download as PDF", tmp.name, file_name="law_chat.pdf")

    elif format_type == "docx":
        doc = Document()
        doc.add_heading("Law Chat Conversation", level=1)
        for msg in messages:
            doc.add_paragraph(f"{msg['role'].capitalize()}: {msg['content']}")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            st.download_button("📘 Download as DOCX", tmp.name, file_name="law_chat.docx")


# ------------------------------
# Main Chatbot UI
# ------------------------------
def show_law_chatbot():
    st.subheader("⚖️ Law Jublie — Constitutional & IPC Legal Assistant")
    coi_df = load_coi_dataset()

    # Export chat
    with st.expander("💾 Export Chat History"):
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📄 Export as PDF"):
                export_chat("pdf")
        with col2:
            if st.button("📘 Export as DOCX"):
                export_chat("docx")

    st.divider()

    if "law_messages" not in st.session_state:
        st.session_state.law_messages = []

    if "show_complaint_form" not in st.session_state:
        st.session_state.show_complaint_form = False

    # Display chat history
    for msg in st.session_state.law_messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # PDF Upload
    pdf_file = st.file_uploader("📂 Upload a law-related PDF", type=["pdf"])
    if pdf_file:
        with st.spinner("🔍 Extracting and analyzing your document..."):
            extracted_text = read_pdf(pdf_file)
            if not extracted_text.strip():
                st.warning("⚠️ Could not extract any readable text from the PDF.")
                return
            if not is_law_related(extracted_text):
                msg = "⚠️ This PDF doesn't seem related to IPC or the Constitution of India."
                st.session_state.law_messages.append({"role": "assistant", "content": msg})
                st.chat_message("assistant").markdown(msg)
            else:
                summary = generate_summary(extracted_text)
                st.session_state.law_messages.append({"role": "assistant", "content": summary})
                st.chat_message("assistant").markdown(summary)

    st.divider()

    # User Input
    col1, col2 = st.columns([8, 2])
    with col1:
        law_query = st.chat_input("💬 Ask a law-related question...", key="law_input")
    with col2:
        if st.button("📢 File Complaint"):
            st.session_state.show_complaint_form = not st.session_state.show_complaint_form

    # Complaint Form
    if st.session_state.show_complaint_form:
        st.subheader("📝 Submit a Complaint")
        with st.form("complaint_form"):
            title = st.text_input("Complaint Title")
            description = st.text_area("Complaint Description")
            submitted = st.form_submit_button("Submit")
            if submitted:
                if title and description:
                    with open("complaints.txt", "a") as f:
                        f.write(f"Title: {title}\nDescription: {description}\n\n")
                    st.success("✅ Complaint submitted successfully.")
                    st.session_state.show_complaint_form = False
                else:
                    st.error("⚠️ Both title and description are required.")

    # Handle Chat Query
    if law_query:
        st.session_state.law_messages.append({"role": "user", "content": law_query})
        with st.chat_message("user"):
            st.markdown(law_query)

        if not is_law_related(law_query):
            msg = "❗ This chatbot only handles Constitution and IPC-related queries."
            st.session_state.law_messages.append({"role": "assistant", "content": msg})
            st.chat_message("assistant").markdown(msg)
            return

        with st.spinner("⚖️ Jublie is analyzing your query..."):
            try:
                system_prompt = (
                    "You are 'Jublie', a senior AI legal advisor specialized in Indian law. "
                    "Provide concise, factual, and verified legal answers based on IPC or the Constitution of India. "
                    "Avoid opinions or assumptions."
                )

                completion = client.chat.completions.create(
                    model=CHAT_MODEL,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": law_query}
                    ],
                    temperature=0.3,
                    max_tokens=1024
                )

                reply = completion.choices[0].message.content.strip()

                # Add verified legal reference
                match = coi_df[
                    coi_df["article"].str.lower().str.contains(law_query.lower(), na=False)
                    | coi_df["title"].str.lower().str.contains(law_query.lower(), na=False)
                ]
                if not match.empty:
                    reply += f"\n\n📘 **Verified Legal Reference:**\n{match.iloc[0]['description']}"

                st.session_state.law_messages.append({"role": "assistant", "content": reply})
                with st.chat_message("assistant"):
                    st.markdown(reply, unsafe_allow_html=True)

            except Exception as e:
                msg = str(e).lower()
                if "rate_limit" in msg:
                    st.warning("⚠️ Groq API rate limit reached. Try again later.")
                elif "invalid_api_key" in msg:
                    st.error("❌ Invalid Groq API key. Check your .env file.")
                elif "model_decommissioned" in msg:
                    st.error("❌ Model decommissioned. Update to a supported model.")
                else:
                    st.error(f"⚠️ Unexpected Groq API error: {e}")


# ------------------------------
# Run Chatbot
# ------------------------------
if __name__ == "__main__":
    show_law_chatbot()
